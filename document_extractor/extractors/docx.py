from pathlib import Path
from typing import Union
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from ..models import DocumentResult

class DOCXExtractor:
    """Extractor for DOCX files using python-docx."""
    
    @staticmethod
    def extract(file_path: Union[str, Path]) -> DocumentResult:
        """Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            DocumentResult containing the extracted text and metadata
        """
        file_path = Path(file_path)
        
        try:
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
                
            if file_path.suffix.lower() != '.docx':
                raise ValueError(f"Not a DOCX file: {file_path}")
            
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_text.append(" | ".join(cells))
            
            # Combine all text with appropriate spacing
            all_text = paragraphs + table_text
            return DocumentResult.from_path(file_path, content="\n\n".join(all_text))
            
        except (FileNotFoundError, ValueError) as e:
            # Pass through common errors with their messages
            return DocumentResult.from_path(file_path, error=str(e))
        except PackageNotFoundError:
            return DocumentResult.from_path(file_path, error="Invalid or corrupted DOCX file")
        except Exception as e:
            return DocumentResult.from_path(file_path, error=f"Unexpected error during DOCX extraction: {e}")